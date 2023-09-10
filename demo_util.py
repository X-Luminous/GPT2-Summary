# -*- coding: utf-8 -*-
# @File:demo_util.py
# @Author:25833
# @Time:2023/5/31

import io
import os
import re
import pandas as pd
import numpy as np
import torch
from tqdm import  tqdm
import time
import datetime
import pdfplumber
import pyofd
import docx

def file_name_walk(file_dir):
    file_path_name_list = [[root, dirs, files] for root, dirs, files in os.walk(file_dir)][0]
    file_path_str = file_path_name_list[0]
    file_name_list = file_path_name_list[2]
    #print(len(file_title_list))
    return file_path_str,file_name_list


def read_pdf(pdf_file):
    content_list = []
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            # print(page.extract_text())
            line_multy = page.extract_text()
            fields = ["%s\n"%(line) for line in line_multy.strip().split("\n")]
            content_list = content_list + fields
    return content_list


def read_ofd(file_dir, encoding_str="gbk"):
    content_list = []
    """
    content_list = []
    with io.open(file_dir,errors="ignore") as f:
        for line in f:
            content_list.append(line)
    """
    """
    content_list = []
    doc = docx.Document(file_dir)  
    idPara_list = doc.paragraphs
    content_list = [dPara.text for dPara in idPara_list]
    """
    #print(pyofd)
    ofdReader = pyofd.OFDReader(file_dir)
    #print(help(ofdReader))
    return content_list


def read_txt(file_dir, encoding_str="utf-8"):
    content_list = []
    with io.open(file_dir, encoding=encoding_str) as f:
        for line in f:
            content_list.append(line)
    return content_list


def read_doc(file_dir):
    doc = docx.Document(file_dir)
    idPara_list = doc.paragraphs
    content_list_p = [dPara.text for dPara in idPara_list]
    content_list = []
    for p in content_list_p:
        #print(p)
        content_list = content_list + p.strip().split("\n")
    return content_list


def read_wps(file_dir, encoding_str="utf-8"):
    content_list = []
    with io.open(file_dir, encoding=encoding_str) as f:
        for line in f:
            content_list.append(line)
    return content_list


def format_data(file_path, encoding_str="utf-8"):
    file_type_count_dict = {}
    file_dict = {}
    file_path_str,file_name_list = file_name_walk(file_path)
    # for file_path in file_path_list:
    #print(file_path_list[0])
    for file_name in tqdm(file_name_list):
        file_path = "%s/%s"%(file_path_str,file_name)
        file_type = file_path.split(".")[-1]
        if file_type in file_type_count_dict:
            file_type_count_dict[file_type] += 1
        else:
            file_type_count_dict[file_type] = 1
        if file_type == "ofd":
            #print(file_path)
            continue
        #if not(file_type == "doc" or file_type == "docx"):
        #    continue

        if file_type == "txt":  # txt
            file_dict[file_name] = read_txt(file_path, encoding_str)
        elif file_type == "doc" or file_type == "docx":  # doc
            file_dict[file_name] = read_doc(file_path)
        elif file_type == "ofd":  # ofd
            file_dict[file_name] = read_ofd(file_path)
        elif file_type == "wps":  # wps
            file_dict[file_name] = read_wps(file_path)
        else:  # pdf
            file_dict[file_name] = read_pdf(file_path)

    print(file_type_count_dict)
    return file_dict


def find_book_list(sentences,i):
    curr_sentences = []
    max_size = len(sentences)
    if i >= 2  and i < max_size-3:
        curr_sentences = sentences[i-2:i+3]
    elif i >= 2  and i >= max_size-3:
        curr_sentences = sentences[i-2:]
    elif i < 2  and i <= max_size-3:
        curr_sentences = sentences[0:i+3]
    else:
        curr_sentences = sentences[0:-1]

    sentence_str = "".join(curr_sentences)
    #print(sentence_str)
    book_name_pattern = r'《([^《》]+)》'
    books_src_list = re.findall(book_name_pattern, sentence_str)
    books_src_list = ["《%s》"%(book) for book in books_src_list]
    books_list = []
    for book_name in  books_src_list:
        fields = book_name.split("\n")
        books_list = books_list + fields
    books_list = list(set(books_list))
    return books_list


if __name__ == '__main__':
    mode_str = "policy"
    base = "E:\\pycharm_script_space\\data_tzb_2023\\%s/"%(mode_str)
    file_content_dict = format_data(base)
    file_name = "内蒙古自治区人民政府关于进一步加强全区城乡规划管理工作的通知.docx"
    print(file_name in file_content_dict.keys())
    print(file_content_dict.get(file_name))
    print(len(file_content_dict.get(file_name)))

